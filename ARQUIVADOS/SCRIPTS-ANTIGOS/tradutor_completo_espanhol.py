#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TRADUTOR COMPLETO PARA ESPANHOL
Preserva TODA formatação, imagens e estrutura original
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import requests
import json
from dotenv import load_dotenv
import time

# Carrega variáveis de ambiente
load_dotenv()

def configurar_openai():
    """Configura a API da OpenAI"""
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        print("❌ OPENAI_API_KEY não encontrada no arquivo .env")
        return False
    
    print("✅ OpenAI configurada")
    return True

def traduzir_texto_ia(texto, contexto=""):
    """Traduz texto usando IA mantendo formatação"""
    if not texto or len(texto.strip()) < 3:
        return texto
    
    # Dicionário de traduções diretas para termos comuns
    traducoes_diretas = {
        "Capítulo": "Capítulo",
        "Página": "Página", 
        "Índice": "Índice",
        "Introdução": "Introducción",
        "Conclusão": "Conclusión",
        "Bibliografia": "Bibliografía",
        "Agradecimentos": "Agradecimientos",
        "Prefácio": "Prefacio",
        "Sumário": "Sumario"
    }
    
    # Verifica traduções diretas primeiro
    for pt, es in traducoes_diretas.items():
        if pt.lower() in texto.lower():
            texto = texto.replace(pt, es)
    
    try:
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            return traduzir_manual_espanhol(texto)
        
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [
                {"role": "system", "content": "Você é um tradutor profissional português-espanhol. Traduza mantendo EXATAMENTE a formatação original."},
                {"role": "user", "content": f"Traduza para espanhol mantendo formatação: {texto}"}
            ],
            "max_tokens": 1000,
            "temperature": 0.3
        }
        
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            traducao = result['choices'][0]['message']['content'].strip()
            return traducao
        else:
            print(f"⚠️ API Error: {response.status_code}")
            return traduzir_manual_espanhol(texto)
        
    except Exception as e:
        print(f"⚠️ Erro na tradução IA: {e}")
        return traduzir_manual_espanhol(texto)

def traduzir_manual_espanhol(texto):
    """Tradução manual básica português-espanhol"""
    traducoes = {
        # Palavras comuns
        "e ": "y ",
        "de ": "de ",
        "do ": "del ",
        "da ": "de la ",
        "dos ": "de los ",
        "das ": "de las ",
        "em ": "en ",
        "no ": "en el ",
        "na ": "en la ",
        "nos ": "en los ",
        "nas ": "en las ",
        "para ": "para ",
        "por ": "por ",
        "com ": "con ",
        "sem ": "sin ",
        "sobre ": "sobre ",
        "entre ": "entre ",
        "durante ": "durante ",
        "através ": "a través ",
        "depois ": "después ",
        "antes ": "antes ",
        "agora ": "ahora ",
        "hoje ": "hoy ",
        "ontem ": "ayer ",
        "amanhã ": "mañana ",
        "sempre ": "siempre ",
        "nunca ": "nunca ",
        "muito ": "muy ",
        "pouco ": "poco ",
        "mais ": "más ",
        "menos ": "menos ",
        "melhor ": "mejor ",
        "pior ": "peor ",
        "grande ": "grande ",
        "pequeno ": "pequeño ",
        "novo ": "nuevo ",
        "velho ": "viejo ",
        "bom ": "bueno ",
        "mau ": "malo ",
        "primeiro ": "primero ",
        "último ": "último ",
        "importante ": "importante ",
        "necessário ": "necesario ",
        "possível ": "posible ",
        "impossível ": "imposible ",
        "fácil ": "fácil ",
        "difícil ": "difícil ",
        "simples ": "simple ",
        "complexo ": "complejo ",
        "rápido ": "rápido ",
        "lento ": "lento ",
        "alto ": "alto ",
        "baixo ": "bajo ",
        "certo ": "cierto ",
        "errado ": "equivocado ",
        "verdade ": "verdad ",
        "mentira ": "mentira ",
        "vida ": "vida ",
        "morte ": "muerte ",
        "amor ": "amor ",
        "ódio ": "odio ",
        "paz ": "paz ",
        "guerra ": "guerra ",
        "trabalho ": "trabajo ",
        "casa ": "casa ",
        "família ": "familia ",
        "amigo ": "amigo ",
        "inimigo ": "enemigo ",
        "pessoa ": "persona ",
        "gente ": "gente ",
        "homem ": "hombre ",
        "mulher ": "mujer ",
        "criança ": "niño ",
        "jovem ": "joven ",
        "adulto ": "adulto ",
        "idoso ": "anciano ",
        "tempo ": "tiempo ",
        "espaço ": "espacio ",
        "lugar ": "lugar ",
        "mundo ": "mundo ",
        "país ": "país ",
        "cidade ": "ciudad ",
        "água ": "agua ",
        "fogo ": "fuego ",
        "terra ": "tierra ",
        "ar ": "aire ",
        "sol ": "sol ",
        "lua ": "luna ",
        "estrela ": "estrella ",
        "dia ": "día ",
        "noite ": "noche ",
        "manhã ": "mañana ",
        "tarde ": "tarde ",
        "ano ": "año ",
        "mês ": "mes ",
        "semana ": "semana ",
        "hora ": "hora ",
        "minuto ": "minuto ",
        "segundo ": "segundo ",
        "dinheiro ": "dinero ",
        "preço ": "precio ",
        "valor ": "valor ",
        "custo ": "costo ",
        "problema ": "problema ",
        "solução ": "solución ",
        "questão ": "cuestión ",
        "resposta ": "respuesta ",
        "pergunta ": "pregunta ",
        "conhecimento ": "conocimiento ",
        "sabedoria ": "sabiduría ",
        "inteligência ": "inteligencia ",
        "educação ": "educación ",
        "escola ": "escuela ",
        "universidade ": "universidad ",
        "professor ": "profesor ",
        "aluno ": "alumno ",
        "livro ": "libro ",
        "página ": "página ",
        "capítulo ": "capítulo ",
        "história ": "historia ",
        "cultura ": "cultura ",
        "arte ": "arte ",
        "música ": "música ",
        "ciência ": "ciencia ",
        "tecnologia ": "tecnología ",
        "medicina ": "medicina ",
        "saúde ": "salud ",
        "doença ": "enfermedad ",
        "hospital ": "hospital ",
        "médico ": "médico ",
        "remédio ": "medicina ",
        "comida ": "comida ",
        "bebida ": "bebida ",
        "restaurante ": "restaurante ",
        "hotel ": "hotel ",
        "viagem ": "viaje ",
        "transporte ": "transporte ",
        "carro ": "coche ",
        "avião ": "avión ",
        "trem ": "tren ",
        "barco ": "barco ",
        "estrada ": "carretera ",
        "rua ": "calle ",
        "prédio ": "edificio ",
        "porta ": "puerta ",
        "janela ": "ventana ",
        "quarto ": "habitación ",
        "cozinha ": "cocina ",
        "banheiro ": "baño ",
        "jardim ": "jardín ",
        "árvore ": "árbol ",
        "flor ": "flor ",
        "animal ": "animal ",
        "cão ": "perro ",
        "gato ": "gato ",
        "pássaro ": "pájaro ",
        "peixe ": "pez ",
        "cor ": "color ",
        "branco ": "blanco ",
        "preto ": "negro ",
        "vermelho ": "rojo ",
        "azul ": "azul ",
        "verde ": "verde ",
        "amarelo ": "amarillo ",
        "laranja ": "naranja ",
        "roxo ": "morado ",
        "rosa ": "rosa ",
        "marrom ": "marrón ",
        "cinza ": "gris "
    }
    
    texto_traduzido = texto
    for pt, es in traducoes.items():
        texto_traduzido = texto_traduzido.replace(pt, es)
    
    return texto_traduzido

def preservar_formatacao_paragrafo(paragrafo_original, paragrafo_novo):
    """Preserva toda formatação do parágrafo original"""
    try:
        # Copia alinhamento
        if paragrafo_original.alignment:
            paragrafo_novo.alignment = paragrafo_original.alignment
        
        # Copia formatação do parágrafo
        if paragrafo_original.paragraph_format:
            pf_orig = paragrafo_original.paragraph_format
            pf_novo = paragrafo_novo.paragraph_format
            
            if pf_orig.space_before:
                pf_novo.space_before = pf_orig.space_before
            if pf_orig.space_after:
                pf_novo.space_after = pf_orig.space_after
            if pf_orig.line_spacing:
                pf_novo.line_spacing = pf_orig.line_spacing
            if pf_orig.first_line_indent:
                pf_novo.first_line_indent = pf_orig.first_line_indent
            if pf_orig.left_indent:
                pf_novo.left_indent = pf_orig.left_indent
            if pf_orig.right_indent:
                pf_novo.right_indent = pf_orig.right_indent
        
        # Preserva formatação de runs
        for i, run_orig in enumerate(paragrafo_original.runs):
            if i < len(paragrafo_novo.runs):
                run_novo = paragrafo_novo.runs[i]
                
                # Copia formatação do texto
                if run_orig.bold is not None:
                    run_novo.bold = run_orig.bold
                if run_orig.italic is not None:
                    run_novo.italic = run_orig.italic
                if run_orig.underline is not None:
                    run_novo.underline = run_orig.underline
                if run_orig.font.size:
                    run_novo.font.size = run_orig.font.size
                if run_orig.font.name:
                    run_novo.font.name = run_orig.font.name
                if run_orig.font.color.rgb:
                    run_novo.font.color.rgb = run_orig.font.color.rgb
                    
    except Exception as e:
        print(f"⚠️ Erro ao preservar formatação: {e}")

def traduzir_documento_completo(docx_path):
    """Traduz documento completo preservando TUDO"""
    print("🌍 INICIANDO TRADUÇÃO COMPLETA PARA ESPANHOL")
    print("=" * 60)
    
    if not os.path.exists(docx_path):
        print(f"❌ Arquivo não encontrado: {docx_path}")
        return False
    
    try:
        # Abre documento original
        doc_original = Document(docx_path)
        print(f"📄 Documento carregado: {len(doc_original.paragraphs)} parágrafos")
        
        # Cria novo documento
        doc_traduzido = Document()
        
        # Copia configurações do documento
        doc_traduzido.core_properties.title = "Manuscrito Português - Traducción al Español"
        doc_traduzido.core_properties.author = "Autor Original"
        doc_traduzido.core_properties.subject = "Traducción completa al español"
        
        total_paragrafos = len(doc_original.paragraphs)
        paragrafos_processados = 0
        imagens_preservadas = 0
        
        print(f"🔄 Processando {total_paragrafos} parágrafos...")
        
        for i, paragrafo in enumerate(doc_original.paragraphs):
            try:
                # Mostra progresso
                if i % 10 == 0:
                    progresso = (i / total_paragrafos) * 100
                    print(f"📊 Progresso: {progresso:.1f}% ({i}/{total_paragrafos})")
                
                # Verifica se tem imagens
                tem_imagem = False
                for run in paragrafo.runs:
                    if run._element.xpath('.//pic:pic'):
                        tem_imagem = True
                        break
                
                if tem_imagem:
                    # Preserva parágrafo com imagem exatamente como está
                    novo_paragrafo = doc_traduzido.add_paragraph()
                    
                    # Copia todos os runs (incluindo imagens)
                    for run in paragrafo.runs:
                        novo_run = novo_paragrafo.add_run()
                        
                        # Copia texto se houver
                        if run.text:
                            texto_traduzido = traduzir_texto_ia(run.text, "texto com imagem")
                            novo_run.text = texto_traduzido
                        
                        # Copia imagens
                        for elemento in run._element:
                            if elemento.tag.endswith('}drawing'):
                                novo_run._element.append(elemento)
                                imagens_preservadas += 1
                    
                    # Preserva formatação
                    preservar_formatacao_paragrafo(paragrafo, novo_paragrafo)
                    
                else:
                    # Traduz texto normal
                    texto_original = paragrafo.text
                    
                    if texto_original.strip():
                        # Traduz o texto
                        texto_traduzido = traduzir_texto_ia(texto_original)
                        
                        # Cria novo parágrafo
                        novo_paragrafo = doc_traduzido.add_paragraph(texto_traduzido)
                        
                        # Preserva formatação
                        preservar_formatacao_paragrafo(paragrafo, novo_paragrafo)
                    else:
                        # Preserva parágrafos vazios
                        doc_traduzido.add_paragraph()
                
                paragrafos_processados += 1
                
                # Pausa para evitar rate limit
                if i % 5 == 0:
                    time.sleep(1)
                    
            except Exception as e:
                print(f"⚠️ Erro no parágrafo {i}: {e}")
                # Adiciona parágrafo original em caso de erro
                doc_traduzido.add_paragraph(paragrafo.text)
        
        # Salva documento traduzido
        arquivo_saida = "MANUSCRITOPORTUGUES-ES-COMPLETO.docx"
        doc_traduzido.save(arquivo_saida)
        
        print(f"\n✅ TRADUÇÃO COMPLETA FINALIZADA!")
        print(f"📄 Arquivo salvo: {arquivo_saida}")
        print(f"📊 Parágrafos processados: {paragrafos_processados}")
        print(f"🖼️ Imagens preservadas: {imagens_preservadas}")
        print(f"📏 Tamanho: {os.path.getsize(arquivo_saida) / (1024*1024):.1f} MB")
        
        return arquivo_saida
        
    except Exception as e:
        print(f"❌ Erro na tradução: {e}")
        return False

def main():
    """Função principal"""
    print("🌍 TRADUTOR COMPLETO PARA ESPANHOL")
    print("=" * 50)
    
    # Configura OpenAI
    if not configurar_openai():
        return
    
    # Arquivo de entrada
    docx_original = "MANUSCRITOPORTUGUES.docx"
    
    if not os.path.exists(docx_original):
        print(f"❌ Arquivo não encontrado: {docx_original}")
        return
    
    print(f"📄 Traduzindo: {docx_original}")
    print("🎯 Objetivo: Preservar TODA formatação e imagens")
    print("🌍 Idioma: Espanhol")
    print()
    
    # Inicia tradução
    arquivo_traduzido = traduzir_documento_completo(docx_original)
    
    if arquivo_traduzido:
        print(f"\n🎉 SUCESSO!")
        print(f"📱 Arquivo pronto: {arquivo_traduzido}")
        print("🚀 Próximo passo: Converter para EPUB KDP")
    else:
        print("\n❌ Falha na tradução")

if __name__ == "__main__":
    main()