#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para criar documento DOCX formatado profissionalmente
do livro "Turma da Aventura: Os Ecos de Constantinopla"
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def add_cover_page(doc):
    """Adiciona página de capa"""
    # Título principal
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("TURMA DA AVENTURA")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("Os Ecos de Constantinopla")
    subtitle_run.font.size = Pt(32)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(139, 0, 0)

    # Linha decorativa
    doc.add_paragraph("\n" * 3)

    # Livro da série
    series = doc.add_paragraph()
    series.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    series_run = series.add_run("Livro 2 da Série Lugares Incríveis")
    series_run.font.size = Pt(14)
    series_run.font.italic = True

    # Espaço
    doc.add_paragraph("\n" * 8)

    # Informação adicional
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info_run = info.add_run("Uma Aventura Através do Tempo")
    info_run.font.size = Pt(12)
    info_run.font.italic = True

    doc.add_page_break()

def add_credits_page(doc):
    """Adiciona página de créditos"""
    doc.add_paragraph("\n" * 5)

    credits = doc.add_paragraph()
    credits.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    credits.add_run("Série: TURMA DA AVENTURA - LUGARES INCRÍVEIS\n\n").font.bold = True
    credits.add_run("Livro 2: Os Ecos de Constantinopla\n\n")
    credits.add_run("Gênero: Aventura Infantil / Fantasia / Ficção Histórica\n")
    credits.add_run("Faixa Etária: 8-12 anos\n\n")
    credits.add_run("Data: Dezembro de 2025\n\n")
    credits.add_run("© 2025 Todos os direitos reservados\n")
    credits.add_run("Este livro faz parte da série Lugares Incríveis\n\n")
    credits.add_run("Livro anterior: Atlântida (Livro 1)\n")
    credits.add_run("Próxima aventura: Os Segredos das Pirâmides (Livro 3)")

    doc.add_page_break()

def add_table_of_contents(doc, chapters):
    """Adiciona índice/sumário"""
    # Título do índice
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_run = toc_title.add_run("SUMÁRIO")
    toc_run.font.size = Pt(20)
    toc_run.font.bold = True

    doc.add_paragraph("\n")

    # Lista de capítulos
    for i, chapter_title in enumerate(chapters, 1):
        toc_item = doc.add_paragraph()
        toc_item.add_run(f"Capítulo {i}: ").font.bold = True
        toc_item.add_run(chapter_title)

    doc.add_page_break()

def process_chapter_content(doc, chapter_num, chapter_title, content):
    """Processa e adiciona conteúdo do capítulo"""
    # Título do capítulo
    chap_title = doc.add_paragraph()
    chap_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = chap_title.add_run(f"CAPÍTULO {chapter_num}")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # Nome do capítulo
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = name.add_run(chapter_title.upper())
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(139, 0, 0)

    doc.add_paragraph()

    # Processar secoes (topicos)
    sections = content.split('---')

    for i, section in enumerate(sections):
        # Tentar inserir imagem
        topic_num = i + 1
        image_filename = f"chapter_{chapter_num}_topic_{topic_num}.png"
        image_path = os.path.join(r"d:\TRAE-PROJETOS\livro1\BIZANTINO\images", image_filename)

        if i > 0:
            doc.add_paragraph()
            separator = doc.add_paragraph()
            separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator.add_run("• • •").font.bold = True
            doc.add_paragraph()

        if os.path.exists(image_path):
            try:
                print(f"  Inserindo imagem: {image_filename}")
                doc.add_picture(image_path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph()
            except Exception as e:
                print(f"  Erro ao inserir imagem {image_filename}: {e}")

        # Processar parágrafos
        paragraphs = section.strip().split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if para.startswith('**') and para.endswith('**'):
                # Subtítulo em negrito
                subtitle_para = doc.add_paragraph()
                subtitle_run = subtitle_para.add_run(para.strip('*'))
                subtitle_run.font.size = Pt(14)
                subtitle_run.font.bold = True
            else:
                # Parágrafo normal
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Inches(0.5)

                # Processar texto com possível formatação
                parts = re.split(r'(\*\*.*?\*\*)', para)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Negrito inline
                        run = p.add_run(part.strip('*'))
                        run.font.bold = True
                    else:
                        p.add_run(part)

    doc.add_page_break()

def add_back_cover(doc):
    """Adiciona contra-capa com sinopse"""
    doc.add_paragraph("\n" * 2)

    # Título da sinopse
    synopsis_title = doc.add_paragraph()
    synopsis_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = synopsis_title.add_run("SOBRE ESTE LIVRO")
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    doc.add_paragraph("\n")

    # Sinopse
    synopsis = doc.add_paragraph()
    synopsis.paragraph_format.line_spacing = 1.5
    synopsis.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    synopsis.add_run(
        "Quando cristais atlantes antigos começam a falhar sob a Hagia Sophia, "
        "a Turma da Aventura deve mergulhar quinze séculos no passado para salvar "
        "uma das maiores maravilhas arquitetônicas do mundo. Em Constantinopla do "
        "ano 537 d.C., Will, Mia, Leo, Sophie, Max, Jimmy e Íris encontram Theodora - "
        "uma órfã esperta das ruas que se tornará guardiã através dos séculos.\n\n"
    )

    synopsis.add_run(
        "Mas forças conspiram para destruir a catedral antes mesmo de ser consagrada. "
        "Em uma corrida contra o tempo através da cidade imperial, enfrentando vilões "
        "complexos e dilemas impossíveis, a turma precisa usar inteligência, coragem "
        "e trabalho em equipe para proteger não apenas um edifício, mas a própria "
        "possibilidade de conexão entre passado e futuro.\n\n"
    )

    synopsis_run = synopsis.add_run(
        "Uma aventura épica de amizade através do tempo, onde cada escolha pode "
        "mudar a história - ou preservá-la para sempre."
    )
    synopsis_run.font.italic = True

    doc.add_paragraph("\n" * 2)

    # Próxima aventura
    next_adv = doc.add_paragraph()
    next_adv.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    next_run = next_adv.add_run("PRÓXIMA AVENTURA")
    next_run.font.size = Pt(14)
    next_run.font.bold = True
    next_run.font.color.rgb = RGBColor(139, 0, 0)

    next_title = doc.add_paragraph()
    next_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    next_title.add_run("Os Segredos das Pirâmides").font.italic = True

    next_desc = doc.add_paragraph()
    next_desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    next_desc.add_run("Egito Antigo aguarda os Restauradores Temporais!")

def main():
    """Função principal"""
    print("Iniciando criação do documento DOCX...")

    # Ler manuscrito
    print("Lendo manuscrito...")
    with open('D:\\TRAE-PROJETOS\\livro1\\BIZANTINO\\LIVROS\\SERIE-LUGARES-INCRIVEIS\\LIVRO-02-CONSTANTINOPLA\\PT\\manuscrito.txt', 'r', encoding='utf-8') as f:
        content = f.read()

    # Criar documento
    doc = Document()

    # Configurar margens (2.5cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Definir fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    print("Adicionando capa...")
    add_cover_page(doc)

    print("Adicionando página de créditos...")
    add_credits_page(doc)

    # Extrair capítulos
    print("Processando capítulos...")
    chapter_pattern = r'CAPÍTULO (\d+):\s*([^\n]+)'
    chapters_matches = list(re.finditer(chapter_pattern, content))

    chapter_titles = []
    for match in chapters_matches:
        chapter_titles.append(match.group(2).strip())

    print(f"Encontrados {len(chapter_titles)} capítulos")

    print("Adicionando índice...")
    add_table_of_contents(doc, chapter_titles)

    # Processar cada capítulo
    for i, match in enumerate(chapters_matches):
        chapter_num = int(match.group(1))
        chapter_title = match.group(2).strip()

        # Encontrar conteúdo do capítulo
        start_pos = match.end()

        # Encontrar fim do capítulo (início do próximo ou fim do arquivo)
        if i < len(chapters_matches) - 1:
            end_pos = chapters_matches[i + 1].start()
        else:
            # Último capítulo - encontrar "**FIM**" ou similar
            fim_match = re.search(r'\*\*FIM\*\*', content[start_pos:])
            if fim_match:
                end_pos = start_pos + fim_match.start()
            else:
                end_pos = len(content)

        chapter_content = content[start_pos:end_pos].strip()

        # Remover linhas de resumo executivo ou notas técnicas se existirem
        chapter_content = re.sub(r'\[Resumo Executivo.*?\]', '', chapter_content, flags=re.DOTALL)
        chapter_content = re.sub(r'\[CONTINUARÁ.*?\]', '', chapter_content, flags=re.DOTALL)
        chapter_content = re.sub(r'\[NOTA TÉCNICA.*?\]', '', chapter_content, flags=re.DOTALL)
        chapter_content = re.sub(r'═+', '', chapter_content)

        print(f"  Processando Capítulo {chapter_num}: {chapter_title}...")
        process_chapter_content(doc, chapter_num, chapter_title, chapter_content)

    print("Adicionando contra-capa...")
    add_back_cover(doc)

    # Salvar documento
    output_path = 'D:\\TRAE-PROJETOS\\livro1\\BIZANTINO\\LIVROS\\SERIE-LUGARES-INCRIVEIS\\LIVRO-02-CONSTANTINOPLA\\PT\\Turma_da_Aventura_Os_Ecos_de_Constantinopla.docx'
    print(f"Salvando documento em: {output_path}")
    doc.save(output_path)

    print("✅ Documento DOCX criado com sucesso!")
    print(f"📄 Arquivo: {output_path}")
    print(f"📚 Capítulos processados: {len(chapter_titles)}")

if __name__ == "__main__":
    main()
