#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
📚 CONVERSOR AMAZON KDP - DOCX PARA PDF E EPUB
🎯 Converte livros para publicação na Amazon
✅ Otimizado para Kindle Direct Publishing
"""

import os
import sys
from docx import Document
from docx2pdf import convert as docx_to_pdf
import subprocess
import tempfile
import shutil
from pathlib import Path

class ConversorAmazonKDP:
    def __init__(self):
        self.qualidade_pdf = "high"
        self.formato_epub = "epub3"
        
    def instalar_dependencias(self):
        """Instala dependências necessárias"""
        print("📦 Verificando dependências...")
        
        dependencias = [
            "docx2pdf",
            "python-docx",
            "ebooklib",
            "lxml",
            "Pillow"
        ]
        
        for dep in dependencias:
            try:
                __import__(dep.replace("-", "_"))
                print(f"  ✅ {dep}")
            except ImportError:
                print(f"  📥 Instalando {dep}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", dep])
                print(f"  ✅ {dep} instalado!")
    
    def converter_para_pdf_libreoffice(self, arquivo_docx, arquivo_pdf):
        """
        Converte DOCX para PDF usando LibreOffice
        """
        try:
            # Tentar usar LibreOffice
            cmd = [
                "soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(arquivo_pdf),
                arquivo_docx
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                # LibreOffice cria o PDF com nome baseado no DOCX
                nome_base = os.path.splitext(os.path.basename(arquivo_docx))[0]
                pdf_gerado = os.path.join(os.path.dirname(arquivo_pdf), f"{nome_base}.pdf")
                
                if os.path.exists(pdf_gerado) and pdf_gerado != arquivo_pdf:
                    shutil.move(pdf_gerado, arquivo_pdf)
                
                return os.path.exists(arquivo_pdf)
            else:
                return False
                
        except Exception:
            return False
    
    def converter_para_pdf_pypdf(self, arquivo_docx, arquivo_pdf):
        """
        Converte DOCX para PDF usando método alternativo
        """
        try:
            from docx import Document
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            
            # Ler documento DOCX
            doc = Document(arquivo_docx)
            
            # Criar PDF
            pdf_doc = SimpleDocTemplate(arquivo_pdf, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Detectar títulos
                    if any(keyword in para.text.lower() for keyword in ['capítulo', 'chapter', 'livro', 'book']):
                        p = Paragraph(para.text, styles['Title'])
                    else:
                        p = Paragraph(para.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            pdf_doc.build(story)
            return True
            
        except Exception as e:
            print(f"    Erro método alternativo: {e}")
            return False
    
    def converter_para_pdf(self, arquivo_docx, arquivo_pdf):
        """
        Converte DOCX para PDF otimizado para Amazon KDP
        """
        print(f"📄 Convertendo para PDF: {os.path.basename(arquivo_docx)}")
        
        # Método 1: docx2pdf
        try:
            docx_to_pdf(arquivo_docx, arquivo_pdf)
            
            if os.path.exists(arquivo_pdf):
                tamanho = os.path.getsize(arquivo_pdf) / (1024*1024)
                print(f"  ✅ PDF criado (docx2pdf): {tamanho:.1f} MB")
                return True
                
        except Exception as e:
            print(f"  ⚠️ docx2pdf falhou: {e}")
        
        # Método 2: LibreOffice
        print(f"  🔄 Tentando LibreOffice...")
        if self.converter_para_pdf_libreoffice(arquivo_docx, arquivo_pdf):
            if os.path.exists(arquivo_pdf):
                tamanho = os.path.getsize(arquivo_pdf) / (1024*1024)
                print(f"  ✅ PDF criado (LibreOffice): {tamanho:.1f} MB")
                return True
        
        # Método 3: ReportLab (instalar se necessário)
        print(f"  🔄 Tentando método alternativo...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"], 
                                capture_output=True)
            
            if self.converter_para_pdf_pypdf(arquivo_docx, arquivo_pdf):
                if os.path.exists(arquivo_pdf):
                    tamanho = os.path.getsize(arquivo_pdf) / (1024*1024)
                    print(f"  ✅ PDF criado (ReportLab): {tamanho:.1f} MB")
                    return True
        except Exception as e:
            print(f"  ⚠️ Método alternativo falhou: {e}")
        
        print(f"  ❌ Todos os métodos de PDF falharam")
        return False
    
    def converter_para_epub_pandoc(self, arquivo_docx, arquivo_epub):
        """
        Converte DOCX para EPUB usando Pandoc (se disponível)
        """
        print(f"📱 Convertendo para EPUB: {os.path.basename(arquivo_docx)}")
        
        try:
            # Verificar se pandoc está disponível
            subprocess.run(["pandoc", "--version"], 
                         capture_output=True, check=True)
            
            # Comando pandoc otimizado para Amazon
            cmd = [
                "pandoc",
                arquivo_docx,
                "-o", arquivo_epub,
                "--epub-version=3",
                "--epub-cover-image=cover.jpg",  # Se houver capa
                "--toc",
                "--toc-depth=2",
                "--epub-metadata=metadata.xml"  # Se houver metadados
            ]
            
            # Executar conversão
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0 and os.path.exists(arquivo_epub):
                tamanho = os.path.getsize(arquivo_epub) / (1024*1024)
                print(f"  ✅ EPUB criado: {tamanho:.1f} MB")
                return True
            else:
                print(f"  ⚠️ Pandoc falhou: {result.stderr}")
                return False
                
        except FileNotFoundError:
            print(f"  ⚠️ Pandoc não encontrado, tentando método alternativo...")
            return False
        except Exception as e:
            print(f"  ❌ Erro na conversão EPUB: {e}")
            return False
    
    def converter_para_epub_manual(self, arquivo_docx, arquivo_epub):
        """
        Converte DOCX para EPUB usando método manual
        """
        print(f"📱 Convertendo para EPUB (método manual): {os.path.basename(arquivo_docx)}")
        
        try:
            from ebooklib import epub
            
            # Ler documento DOCX
            doc = Document(arquivo_docx)
            
            # Criar livro EPUB
            book = epub.EpubBook()
            
            # Metadados básicos
            nome_arquivo = os.path.splitext(os.path.basename(arquivo_docx))[0]
            book.set_identifier(f'id_{nome_arquivo}')
            book.set_title(nome_arquivo.replace('-', ' ').title())
            book.set_language('en' if 'EN' in arquivo_docx else 'es')
            book.add_author('Turma da Aventura')
            
            # Converter parágrafos para HTML
            html_content = "<html><head><title>Chapter</title></head><body>"
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Detectar títulos
                    if any(keyword in para.text.lower() for keyword in ['capítulo', 'chapter', 'livro', 'book']):
                        html_content += f"<h1>{para.text}</h1>"
                    else:
                        html_content += f"<p>{para.text}</p>"
            
            html_content += "</body></html>"
            
            # Criar capítulo
            c1 = epub.EpubHtml(title='Story', file_name='story.xhtml', lang='en')
            c1.content = html_content
            
            # Adicionar capítulo ao livro
            book.add_item(c1)
            
            # Criar índice
            book.toc = (epub.Link("story.xhtml", "Story", "story"),)
            
            # Adicionar navegação
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Definir ordem de leitura
            book.spine = ['nav', c1]
            
            # Salvar EPUB
            epub.write_epub(arquivo_epub, book, {})
            
            if os.path.exists(arquivo_epub):
                tamanho = os.path.getsize(arquivo_epub) / (1024*1024)
                print(f"  ✅ EPUB criado: {tamanho:.1f} MB")
                return True
            else:
                return False
                
        except Exception as e:
            print(f"  ❌ Erro na conversão EPUB manual: {e}")
            return False
    
    def converter_arquivo(self, arquivo_docx):
        """
        Converte um arquivo DOCX para PDF e EPUB
        """
        if not os.path.exists(arquivo_docx):
            print(f"❌ Arquivo não encontrado: {arquivo_docx}")
            return False, False
        
        # Definir nomes dos arquivos de saída
        base_name = os.path.splitext(arquivo_docx)[0]
        arquivo_pdf = f"{base_name}.pdf"
        arquivo_epub = f"{base_name}.epub"
        
        print(f"\n🚀 CONVERTENDO: {os.path.basename(arquivo_docx)}")
        print(f"📁 Diretório: {os.path.dirname(arquivo_docx)}")
        
        # Converter para PDF
        pdf_sucesso = self.converter_para_pdf(arquivo_docx, arquivo_pdf)
        
        # Converter para EPUB (tentar pandoc primeiro, depois manual)
        epub_sucesso = self.converter_para_epub_pandoc(arquivo_docx, arquivo_epub)
        if not epub_sucesso:
            epub_sucesso = self.converter_para_epub_manual(arquivo_docx, arquivo_epub)
        
        return pdf_sucesso, epub_sucesso
    
    def processar_arquivos(self, arquivos):
        """
        Processa lista de arquivos
        """
        print("=" * 70)
        print("📚 CONVERSOR AMAZON KDP")
        print("🎯 DOCX → PDF + EPUB")
        print("=" * 70)
        
        # Instalar dependências
        self.instalar_dependencias()
        
        resultados = []
        
        for arquivo in arquivos:
            if os.path.exists(arquivo):
                pdf_ok, epub_ok = self.converter_arquivo(arquivo)
                resultados.append({
                    'arquivo': arquivo,
                    'pdf': pdf_ok,
                    'epub': epub_ok
                })
            else:
                print(f"⚠️ Arquivo não encontrado: {arquivo}")
                resultados.append({
                    'arquivo': arquivo,
                    'pdf': False,
                    'epub': False
                })
        
        # Relatório final
        print(f"\n🎉 CONVERSÃO FINALIZADA!")
        print("=" * 50)
        
        for resultado in resultados:
            nome = os.path.basename(resultado['arquivo'])
            pdf_status = "✅" if resultado['pdf'] else "❌"
            epub_status = "✅" if resultado['epub'] else "❌"
            print(f"{nome}")
            print(f"  PDF:  {pdf_status}")
            print(f"  EPUB: {epub_status}")
            print()
        
        return resultados

def main():
    """Função principal"""
    conversor = ConversorAmazonKDP()
    
    # Arquivos para converter
    arquivos = [
        r"c:\Users\xberi\Documents\trae_projects\154\LIVROS\LIVRO1\MANUSCRITOPORTUGUES-EN-GRATUITO.docx",
        r"c:\Users\xberi\Documents\trae_projects\154\LIVROS\LIVRO1\MANUSCRITOPORTUGUES-ES-GRATUITO.docx"
    ]
    
    # Processar conversões
    resultados = conversor.processar_arquivos(arquivos)
    
    # Verificar se todos foram convertidos
    total_pdf = sum(1 for r in resultados if r['pdf'])
    total_epub = sum(1 for r in resultados if r['epub'])
    
    print(f"📊 RESUMO:")
    print(f"  PDFs criados: {total_pdf}/{len(arquivos)}")
    print(f"  EPUBs criados: {total_epub}/{len(arquivos)}")
    
    if total_pdf == len(arquivos) and total_epub == len(arquivos):
        print(f"🎉 SUCESSO TOTAL! Arquivos prontos para Amazon KDP!")
    else:
        print(f"⚠️ Algumas conversões falharam. Verifique os logs acima.")

if __name__ == "__main__":
    main()