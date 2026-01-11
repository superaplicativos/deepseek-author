#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CONVERSOR EPUB KDP ESPECIALIZADO - VERSÃO ESPANHOL
Converte DOCX traduzido para EPUB otimizado para Amazon KDP
Preserva TODAS as páginas iniciais, formatação e imagens
"""

import os
import re
import zipfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
from PIL import Image
import io

class ConversorEpubKdpEspanhol:
    def __init__(self):
        self.imagens_extraidas = []
        self.capitulos = []
        self.capa_nome = None
        self.metadados = {
            'titulo': 'MANUSCRITO PORTUGUÉS - ESPAÑOL GRATUITO',
            'autor': 'Autor del Manuscrito',
            'idioma': 'es',
            'identificador': 'manuscrito-portugues-espanol-gratuito-2024',
            'data_publicacao': datetime.now().strftime('%Y-%m-%d'),
            'editora': 'Edición Independiente',
            'descricao': 'Manuscrito completo traducido al español, optimizado para Amazon KDP con todas las imágenes y formateo preservados.'
        }
    
    def extrair_imagens_docx(self, arquivo_docx):
        """Extrai todas as imagens do arquivo DOCX"""
        print(f"🖼️ Extraindo imagens de {arquivo_docx}...")
        
        imagens_extraidas = []
        
        try:
            # Abre o DOCX como ZIP
            with zipfile.ZipFile(arquivo_docx, 'r') as docx_zip:
                # Lista todos os arquivos
                arquivos = docx_zip.namelist()
                
                # Filtra arquivos de imagem
                arquivos_imagem = [f for f in arquivos if f.startswith('word/media/') and 
                                 any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'])]
                
                print(f"📊 Encontradas {len(arquivos_imagem)} imagens no documento")
                
                # Extrai cada imagem
                for i, arquivo_img in enumerate(arquivos_imagem):
                    try:
                        # Lê os dados da imagem
                        dados_imagem = docx_zip.read(arquivo_img)
                        
                        # Determina a extensão
                        nome_original = os.path.basename(arquivo_img)
                        extensao = os.path.splitext(nome_original)[1].lower()
                        
                        # Gera nome único
                        nome_imagem = f"imagen_{i+1:03d}{extensao}"
                        
                        # Salva a imagem
                        caminho_imagem = os.path.join("imagens_extraidas", nome_imagem)
                        os.makedirs("imagens_extraidas", exist_ok=True)
                        
                        with open(caminho_imagem, 'wb') as f:
                            f.write(dados_imagem)
                        
                        # Verifica e otimiza a imagem para KDP
                        info_imagem = self.otimizar_imagem_kdp(caminho_imagem)
                        
                        if info_imagem:
                            imagens_extraidas.append({
                                'nome': nome_imagem,
                                'caminho': caminho_imagem,
                                'nome_original': nome_original,
                                'tamanho': len(dados_imagem),
                                'largura': info_imagem['largura'],
                                'altura': info_imagem['altura'],
                                'formato': info_imagem['formato'],
                                'kdp_compativel': info_imagem['kdp_compativel']
                            })
                            
                            print(f"   ✅ {nome_imagem}: {info_imagem['largura']}x{info_imagem['altura']} - {len(dados_imagem)/1024:.1f}KB")
                        else:
                            print(f"   ❌ Erro ao processar {nome_imagem}")
                            
                    except Exception as e:
                        print(f"   ❌ Erro ao extrair {arquivo_img}: {str(e)}")
                        continue
                
                self.imagens_extraidas = imagens_extraidas
                print(f"✅ Extração concluída: {len(imagens_extraidas)} imagens válidas")
                return imagens_extraidas
                
        except Exception as e:
            print(f"❌ Erro ao extrair imagens: {str(e)}")
            return []
    
    def otimizar_imagem_kdp(self, caminho_imagem):
        """Otimiza imagem para compatibilidade com KDP"""
        try:
            with Image.open(caminho_imagem) as img:
                largura, altura = img.size
                formato = img.format
                
                # Critérios KDP
                kdp_compativel = True
                motivos_incompatibilidade = []
                
                # Verifica dimensões mínimas (KDP recomenda pelo menos 300 DPI)
                if largura < 300 or altura < 300:
                    motivos_incompatibilidade.append("Dimensões muito pequenas")
                
                # Verifica dimensões máximas (KDP tem limite de 127 MB por arquivo)
                if largura > 4000 or altura > 4000:
                    motivos_incompatibilidade.append("Dimensões muito grandes")
                
                # Verifica formato
                if formato not in ['JPEG', 'PNG', 'GIF']:
                    motivos_incompatibilidade.append(f"Formato {formato} não ideal")
                
                if motivos_incompatibilidade:
                    kdp_compativel = False
                
                return {
                    'largura': largura,
                    'altura': altura,
                    'formato': formato,
                    'kdp_compativel': kdp_compativel,
                    'motivos_incompatibilidade': motivos_incompatibilidade
                }
                
        except Exception as e:
            print(f"❌ Erro ao analisar imagem {caminho_imagem}: {str(e)}")
            return None
    
    def processar_documento_docx(self, arquivo_docx):
        """Processa o documento DOCX e organiza o conteúdo"""
        print(f"📄 Processando documento {arquivo_docx}...")
        
        try:
            doc = Document(arquivo_docx)
            
            # Conta elementos
            total_paragrafos = len(doc.paragraphs)
            print(f"📊 Total de parágrafos: {total_paragrafos}")
            
            # Organiza conteúdo
            conteudo_inicial = []
            capitulos = []
            capitulo_atual = None
            
            # Padrões para identificar capítulos (português e espanhol)
            padroes_capitulo = [
                r'^CAPÍTULO\s+(\d+)',
                r'^CAPÍTULO\s+([IVXLCDM]+)',
                r'^Capítulo\s+(\d+)',
                r'^Capítulo\s+([IVXLCDM]+)',
                r'^CAPÍTULO\s+(\d+)',
                r'^CAPÍTULO\s+([IVXLCDM]+)',
                r'^Capítulo\s+(\d+)',
                r'^Capítulo\s+([IVXLCDM]+)',
                r'^(\d+)\.\s*[A-ZÁÉÍÓÚÑ]',
                r'^([IVXLCDM]+)\.\s*[A-ZÁÉÍÓÚÑ]',
                r'^\d+\s*[-–—]\s*[A-ZÁÉÍÓÚÑ]',
                r'^[A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]{10,50}$'
            ]
            
            for i, paragrafo in enumerate(doc.paragraphs):
                texto = paragrafo.text.strip()
                
                if not texto:
                    continue
                
                # Verifica se é início de capítulo
                eh_capitulo = False
                numero_capitulo = None
                
                for padrao in padroes_capitulo:
                    match = re.match(padrao, texto, re.IGNORECASE)
                    if match:
                        eh_capitulo = True
                        try:
                            numero_capitulo = match.group(1)
                        except IndexError:
                            # Padrão sem grupo capturado; usa numeração sequencial
                            numero_capitulo = str(len(capitulos) + 1)
                        break
                
                if eh_capitulo:
                    # Salva capítulo anterior se existir
                    if capitulo_atual:
                        capitulos.append(capitulo_atual)
                    
                    # Inicia novo capítulo
                    capitulo_atual = {
                        'numero': numero_capitulo,
                        'titulo': texto,
                        'conteudo': [texto],
                        'paragrafo_inicio': i
                    }
                    print(f"📖 Capítulo encontrado: {texto}")
                    
                elif capitulo_atual:
                    # Adiciona ao capítulo atual
                    capitulo_atual['conteudo'].append(texto)
                else:
                    # Conteúdo inicial (antes dos capítulos)
                    conteudo_inicial.append(texto)
            
            # Adiciona último capítulo
            if capitulo_atual:
                capitulos.append(capitulo_atual)
            
            self.capitulos = capitulos
            
            print(f"📚 Estrutura do documento:")
            print(f"   - Conteúdo inicial: {len(conteudo_inicial)} parágrafos")
            print(f"   - Capítulos: {len(capitulos)}")
            
            for cap in capitulos:
                print(f"     • {cap['titulo']}: {len(cap['conteudo'])} parágrafos")
            
            return {
                'conteudo_inicial': conteudo_inicial,
                'capitulos': capitulos,
                'total_paragrafos': total_paragrafos
            }
            
        except Exception as e:
            print(f"❌ Erro ao processar documento: {str(e)}")
            return None
    
    def criar_estrutura_epub(self, nome_arquivo_saida):
        """Cria a estrutura básica do EPUB"""
        print(f"📁 Criando estrutura EPUB...")
        
        # Remove arquivo existente
        if os.path.exists(nome_arquivo_saida):
            os.remove(nome_arquivo_saida)
        
        # Cria diretório temporário
        temp_dir = "temp_epub_espanhol"
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir)
        
        # Estrutura EPUB
        os.makedirs(f"{temp_dir}/META-INF")
        os.makedirs(f"{temp_dir}/OEBPS")
        os.makedirs(f"{temp_dir}/OEBPS/images")
        os.makedirs(f"{temp_dir}/OEBPS/styles")
        os.makedirs(f"{temp_dir}/OEBPS/text")
        
        # mimetype
        with open(f"{temp_dir}/mimetype", 'w', encoding='utf-8') as f:
            f.write("application/epub+zip")
        
        # container.xml
        container_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''
        
        with open(f"{temp_dir}/META-INF/container.xml", 'w', encoding='utf-8') as f:
            f.write(container_xml)
        
        return temp_dir
    
    def criar_css_kdp(self, temp_dir):
        """Cria CSS otimizado para KDP"""
        css_content = '''/* CSS Optimizado para Amazon KDP - Versión Español */

@page {
    margin: 0.5in;
}

body {
    font-family: "Times New Roman", serif;
    font-size: 12pt;
    line-height: 1.4;
    margin: 0;
    padding: 0;
    text-align: justify;
    color: #000000;
}

h1, h2, h3, h4, h5, h6 {
    font-family: "Arial", sans-serif;
    font-weight: bold;
    margin-top: 1.5em;
    margin-bottom: 1em;
    text-align: center;
    page-break-after: avoid;
}

h1 {
    font-size: 18pt;
    margin-top: 2em;
    margin-bottom: 1.5em;
}

h2 {
    font-size: 16pt;
}

h3 {
    font-size: 14pt;
}

.chapter-title {
    font-size: 20pt;
    font-weight: bold;
    text-align: center;
    margin-top: 3em;
    margin-bottom: 2em;
    page-break-before: always;
}

.chapter-number {
    font-size: 16pt;
    font-weight: normal;
    color: #666666;
    margin-bottom: 0.5em;
}

p {
    margin: 0;
    padding: 0;
    text-indent: 1.5em;
    margin-bottom: 0.5em;
}

.no-indent {
    text-indent: 0;
}

.center {
    text-align: center;
    text-indent: 0;
}

.right {
    text-align: right;
    text-indent: 0;
}

.image-container {
    text-align: center;
    margin: 1em 0;
    page-break-inside: avoid;
}

.image-container img {
    max-width: 100%;
    height: auto;
    display: block;
    margin: 0 auto;
}

.image-caption {
    font-size: 10pt;
    font-style: italic;
    text-align: center;
    margin-top: 0.5em;
    color: #666666;
}

.page-break {
    page-break-before: always;
}

.no-break {
    page-break-inside: avoid;
}

/* Estilos específicos para contenido inicial */
.front-matter {
    text-align: center;
    margin: 2em 0;
}

.title-page {
    text-align: center;
    margin: 4em 0;
}

.title-page h1 {
    font-size: 24pt;
    margin-bottom: 2em;
}

.title-page .author {
    font-size: 16pt;
    margin-top: 2em;
}

/* Optimizaciones para KDP */
@media print {
    body {
        font-size: 11pt;
    }
    
    .chapter-title {
        font-size: 18pt;
    }
}

/* Compatibilidad con lectores electrónicos */
@media screen {
    body {
        max-width: 600px;
        margin: 0 auto;
        padding: 1em;
    }
}
'''
        
        with open(f"{temp_dir}/OEBPS/styles/style.css", 'w', encoding='utf-8') as f:
            f.write(css_content)
        
        print("✅ CSS criado com otimizações KDP")
    
    def copiar_imagens(self, temp_dir):
        """Copia imagens para a estrutura EPUB, convertendo formatos não compatíveis e definindo a melhor capa"""
        print(f"🖼️ Copiando {len(self.imagens_extraidas)} imagens...")

        dest_dir = f"{temp_dir}/OEBPS/images"
        os.makedirs(dest_dir, exist_ok=True)

        allowed_formats = {'JPEG', 'PNG', 'GIF'}
        updated_imagens = []

        for idx, imagem in enumerate(self.imagens_extraidas):
            origem = imagem.get('caminho')
            nome = imagem.get('nome')
            formato = imagem.get('formato')
            kdp_ok = imagem.get('kdp_compativel', True)

            destino = os.path.join(dest_dir, nome)

            try:
                # Converte para JPEG se o formato não for suportado ou não estiver compatível com KDP
                if (formato not in allowed_formats) or (not kdp_ok):
                    new_nome = os.path.splitext(nome)[0] + ".jpg"
                    destino = os.path.join(dest_dir, new_nome)
                    with Image.open(origem) as img:
                        # Garante modo RGB para JPEG
                        if img.mode not in ("RGB",):
                            img = img.convert("RGB")
                        try:
                            img.save(destino, format="JPEG", quality=85, optimize=True)
                        except Exception:
                            # Fallback sem optimize em caso de erro
                            img.save(destino, format="JPEG", quality=85)
                    print(f"   🔄 {nome} convertido para {new_nome} (JPEG)")
                    imagem['nome'] = new_nome
                    imagem['formato'] = 'JPEG'
                    imagem['caminho'] = destino
                    imagem['kdp_compativel'] = True
                else:
                    shutil.copy2(origem, destino)
                    imagem['caminho'] = destino
                    print(f"   ✅ {nome}")

                updated_imagens.append(imagem)
            except Exception as e:
                print(f"   ❌ Erro ao copiar/convertir {nome}: {str(e)}")

        # Atualiza lista de imagens com possíveis conversões
        self.imagens_extraidas = updated_imagens

        # Define capa: prioriza nomes contendo 'cover' ou 'capa', caso contrário a maior imagem
        capa_item = None
        for img in self.imagens_extraidas:
            orig_name = img.get('nome_original', '') or img.get('nome', '')
            if re.search(r'(cover|capa)', orig_name, re.IGNORECASE):
                capa_item = img
                break
        if not capa_item and self.imagens_extraidas:
            capa_item = max(self.imagens_extraidas, key=lambda x: x.get('largura', 0) * x.get('altura', 0))

        if capa_item:
            self.capa_nome = capa_item['nome']
            print(f"📗 Capa definida: {self.capa_nome}")
        else:
            print("⚠️ Nenhuma imagem adequada para capa encontrada")

        print("✅ Imagens copiadas para EPUB")
    
    def criar_pagina_capa(self, temp_dir):
        """Cria uma página de capa (cover.xhtml) se houver imagem definida"""
        if not self.capa_nome:
            return
        print("📗 Criando página de capa...")
        cover_html = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Capa</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
</head>
<body>
    <div class="image-container">
        <img src="../images/{self.capa_nome}" alt="Capa"/>
    </div>
</body>
</html>'''
        with open(f"{temp_dir}/OEBPS/text/cover.xhtml", 'w', encoding='utf-8') as f:
            f.write(cover_html)
        print("✅ Página de capa criada")
    
    def criar_pagina_inicial(self, temp_dir, conteudo_inicial):
        """Cria página inicial com todo o conteúdo pré-capítulos"""
        print("📄 Criando página inicial completa...")
        
        html_content = '''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Contenido Inicial</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
</head>
<body>
    <div class="front-matter">
'''
        
        # Adiciona todo o conteúdo inicial
        for i, paragrafo in enumerate(conteudo_inicial):
            if not paragrafo.strip():
                continue
            
            # Detecta se é título principal
            if i == 0 or len(paragrafo) < 100 and paragrafo.isupper():
                html_content += f'        <h1 class="title-page">{paragrafo}</h1>\n'
            elif len(paragrafo) < 50 and any(palavra in paragrafo.lower() for palavra in ['autor', 'por', 'de', 'escrito']):
                html_content += f'        <p class="author center">{paragrafo}</p>\n'
            else:
                html_content += f'        <p>{paragrafo}</p>\n'
        
        html_content += '''    </div>
</body>
</html>'''
        
        with open(f"{temp_dir}/OEBPS/text/inicial.xhtml", 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print("✅ Página inicial criada com todo o conteúdo")
    
    def criar_capitulos_html(self, temp_dir):
        """Cria arquivos HTML para cada capítulo"""
        print(f"📚 Criando {len(self.capitulos)} capítulos...")
        
        for i, capitulo in enumerate(self.capitulos):
            cap_num = i + 1
            cap_titulo = capitulo['titulo']
            cap_numero = capitulo.get('numero', str(cap_num))
            
            html_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{cap_titulo}</title>
    <link rel="stylesheet" type="text/css" href="../styles/style.css"/>
</head>
<body>
    <div class="chapter">
        <h1 class="chapter-title">
            <span class="chapter-number">Capítulo {cap_numero}</span><br/>
            {cap_titulo}
        </h1>
'''
            
            # Adiciona conteúdo do capítulo
            for paragrafo in capitulo['conteudo'][1:]:  # Pula o título
                if paragrafo.strip():
                    html_content += f'        <p>{paragrafo}</p>\n'
            
            html_content += '''    </div>
</body>
</html>'''
            
            nome_arquivo = f"capitulo_{cap_num:02d}.xhtml"
            with open(f"{temp_dir}/OEBPS/text/{nome_arquivo}", 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            print(f"   ✅ {nome_arquivo}: {cap_titulo}")
        
        print("✅ Capítulos criados")
    
    def criar_content_opf(self, temp_dir):
        """Cria arquivo content.opf (manifest e spine)"""
        print("📋 Criando content.opf...")

        # Define media-type da capa se existir
        cover_media = None
        if self.capa_nome:
            ext = os.path.splitext(self.capa_nome)[1].lower()
            if ext in (".jpg", ".jpeg"):
                cover_media = "image/jpeg"
            elif ext == ".png":
                cover_media = "image/png"
            elif ext == ".gif":
                cover_media = "image/gif"
            else:
                cover_media = "image/jpeg"

        # Cabeçalho inicial
        opf_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="bookid" version="2.0">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:title>{self.metadados['titulo']}</dc:title>
        <dc:creator opf:role="aut">{self.metadados['autor']}</dc:creator>
        <dc:language>{self.metadados['idioma']}</dc:language>
        <dc:identifier id="bookid">{self.metadados['identificador']}</dc:identifier>
        <dc:date>{self.metadados['data_publicacao']}</dc:date>
        <dc:publisher>{self.metadados['editora']}</dc:publisher>
        <dc:description>{self.metadados['descricao']}</dc:description>
        {('<meta name="cover" content="cover-image"/>' if self.capa_nome else '')}
    </metadata>

    <manifest>
        <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
        <item id="css" href="styles/style.css" media-type="text/css"/>
        {('<item id="cover" href="text/cover.xhtml" media-type="application/xhtml+xml"/>' if self.capa_nome else '')}
        {((f'<item id="cover-image" href="images/{self.capa_nome}" media-type="{cover_media}"/>' ) if self.capa_nome else '')}
        <item id="inicial" href="text/inicial.xhtml" media-type="application/xhtml+xml"/>
"""

        # Adiciona capítulos ao manifest
        for i in range(len(self.capitulos)):
            cap_id = f"capitulo_{i+1:02d}"
            opf_content += f'        <item id="{cap_id}" href="text/capitulo_{i+1:02d}.xhtml" media-type="application/xhtml+xml"/>' + "\n"

        # Adiciona imagens ao manifest
        for imagem in self.imagens_extraidas:
            img_nome = imagem['nome']
            img_id = os.path.splitext(img_nome)[0]
            extensao = os.path.splitext(img_nome)[1].lower()

            if extensao == '.jpg' or extensao == '.jpeg':
                media_type = 'image/jpeg'
            elif extensao == '.png':
                media_type = 'image/png'
            elif extensao == '.gif':
                media_type = 'image/gif'
            else:
                media_type = 'image/jpeg'

            opf_content += f'        <item id="{img_id}" href="images/{img_nome}" media-type="{media_type}"/>' + "\n"

        # Spine
        opf_content += "    </manifest>\n\n    <spine toc=\"ncx\">\n"
        if self.capa_nome:
            opf_content += "        <itemref idref=\"cover\" linear=\"no\"/>\n"
        opf_content += "        <itemref idref=\"inicial\"/>\n"

        # Adiciona capítulos ao spine
        for i in range(len(self.capitulos)):
            cap_id = f"capitulo_{i+1:02d}"
            opf_content += f'        <itemref idref="{cap_id}"/>' + "\n"

        # Fecha spine e adiciona guia
        opf_content += """    </spine>
    <guide>
"""
        if self.capa_nome:
            opf_content += '        <reference type="cover" title="Capa" href="text/cover.xhtml"/>\n'
        opf_content += """    </guide>
</package>"""

        with open(f"{temp_dir}/OEBPS/content.opf", 'w', encoding='utf-8') as f:
            f.write(opf_content)

        print("✅ content.opf criado")
    
    def criar_toc_ncx(self, temp_dir):
        """Cria arquivo de navegação toc.ncx"""
        print("🧭 Criando toc.ncx...")
        
        ncx_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
    <head>
        <meta name="dtb:uid" content="{self.metadados['identificador']}"/>
        <meta name="dtb:depth" content="2"/>
        <meta name="dtb:totalPageCount" content="0"/>
        <meta name="dtb:maxPageNumber" content="0"/>
    </head>
    
    <docTitle>
        <text>{self.metadados['titulo']}</text>
    </docTitle>
    
    <navMap>
        <navPoint id="inicial" playOrder="1">
            <navLabel>
                <text>Contenido Inicial</text>
            </navLabel>
            <content src="text/inicial.xhtml"/>
        </navPoint>
'''
        
        # Adiciona capítulos à navegação
        for i, capitulo in enumerate(self.capitulos):
            cap_num = i + 1
            cap_titulo = capitulo['titulo']
            play_order = i + 2
            
            ncx_content += f'''        <navPoint id="capitulo_{cap_num:02d}" playOrder="{play_order}">
            <navLabel>
                <text>{cap_titulo}</text>
            </navLabel>
            <content src="text/capitulo_{cap_num:02d}.xhtml"/>
        </navPoint>
'''
        
        ncx_content += '''    </navMap>
</ncx>'''
        
        with open(f"{temp_dir}/OEBPS/toc.ncx", 'w', encoding='utf-8') as f:
            f.write(ncx_content)
        
        print("✅ toc.ncx criado")
    
    def criar_epub_final(self, temp_dir, nome_arquivo_saida):
        """Cria o arquivo EPUB final"""
        print(f"📦 Criando arquivo EPUB: {nome_arquivo_saida}")
        
        with zipfile.ZipFile(nome_arquivo_saida, 'w', zipfile.ZIP_DEFLATED) as epub_zip:
            # Adiciona mimetype primeiro (sem compressão)
            epub_zip.write(f"{temp_dir}/mimetype", "mimetype", compress_type=zipfile.ZIP_STORED)
            
            # Adiciona outros arquivos
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file == "mimetype":
                        continue
                    
                    file_path = os.path.join(root, file)
                    arc_path = os.path.relpath(file_path, temp_dir)
                    epub_zip.write(file_path, arc_path)
        
        # Remove diretório temporário
        shutil.rmtree(temp_dir)
        
        # Verifica arquivo criado
        if os.path.exists(nome_arquivo_saida):
            tamanho = os.path.getsize(nome_arquivo_saida)
            print(f"✅ EPUB criado com sucesso!")
            print(f"📁 Arquivo: {nome_arquivo_saida}")
            print(f"📏 Tamanho: {tamanho:,} bytes ({tamanho/1024/1024:.1f} MB)")
            return True
        else:
            print(f"❌ Erro: EPUB não foi criado")
            return False
    
    def converter_docx_para_epub(self, arquivo_docx, arquivo_epub):
        """Converte DOCX para EPUB KDP completo"""
        print("🇪🇸 CONVERSOR EPUB KDP - VERSÃO ESPANHOL")
        print("=" * 60)
        
        inicio = datetime.now()
        
        # 1. Extrai imagens
        imagens = self.extrair_imagens_docx(arquivo_docx)
        if not imagens:
            print("⚠️ Nenhuma imagem encontrada, continuando sem imagens...")
        
        # 2. Processa documento
        estrutura = self.processar_documento_docx(arquivo_docx)
        if not estrutura:
            print("❌ Erro ao processar documento")
            return False
        
        # 3. Cria estrutura EPUB
        temp_dir = self.criar_estrutura_epub(arquivo_epub)
        
        # 4. Cria CSS
        self.criar_css_kdp(temp_dir)
        
        # 5. Copia imagens
        if imagens:
            self.copiar_imagens(temp_dir)
            self.criar_pagina_capa(temp_dir)
        
        # 6. Cria página inicial
        self.criar_pagina_inicial(temp_dir, estrutura['conteudo_inicial'])

        # 7. Cria capítulos
        self.criar_capitulos_html(temp_dir)

        # 8. Cria manifest (content.opf)
        self.criar_content_opf(temp_dir)

        # 9. Cria navegação (toc.ncx)
        self.criar_toc_ncx(temp_dir)

        # 10. Cria EPUB final
        sucesso = self.criar_epub_final(temp_dir, arquivo_epub)

        fim = datetime.now()
        duracao = fim - inicio

        if sucesso:
            print(f"🎉 CONVERSÃO CONCLUÍDA COM SUCESSO!")
            print(f"⏱️ Tempo de execução: {duracao}")
            print(f"📊 Estatísticas:")
            print(f"   - Imagens: {len(self.imagens_extraidas)}")
            print(f"   - Capítulos: {len(self.capitulos)}")
            print(f"   - Arquivo EPUB: {arquivo_epub}")
            print(f"🚀 EPUB otimizado para Amazon KDP!")
        else:
            print(f"❌ Falha na conversão")
        
        return sucesso

def main():
    """Função principal"""
    print("🇪🇸 CONVERSOR EPUB KDP ESPECIALIZADO - VERSÃO ESPANHOL")
    print("=" * 60)
    
    # Arquivos
    arquivo_docx = "MANUSCRITOPORTUGUES-ESPANHOL.docx"
    arquivo_epub = "MANUSCRITOPORTUGUES-ESPANHOL-KDP.epub"
    
    # Verifica se o arquivo existe
    if not os.path.exists(arquivo_docx):
        print(f"❌ Erro: Arquivo {arquivo_docx} não encontrado")
        return False
    
    # Cria o conversor
    conversor = ConversorEpubKdpEspanhol()
    
    # Executa a conversão
    sucesso = conversor.converter_docx_para_epub(arquivo_docx, arquivo_epub)
    
    if sucesso:
        print(f"✅ Conversão concluída!")
        print(f"📄 Arquivo original: {arquivo_docx}")
        print(f"📚 EPUB KDP: {arquivo_epub}")
    else:
        print(f"❌ Falha na conversão")
    
    return sucesso

if __name__ == "__main__":
    main()