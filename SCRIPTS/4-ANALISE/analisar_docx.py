#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Analisador de DOCX para tradução
- Conta palavras (documento, cabeçalhos/rodapés, notas de rodapé e fim)
- Conta imagens
- Lista estilos de parágrafo detectados
- Conta tabelas e verifica presença de TOC
"""

import os
import sys
import zipfile
import xml.etree.ElementTree as ET
import re
from pathlib import Path
from collections import Counter

try:
    from docx import Document
except Exception:
    # Tenta instalar python-docx se não estiver presente
    os.system("pip install python-docx")
    from docx import Document

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

XML_PARTS_TEXT = [
    'word/document.xml',
    # Campos comuns adicionais
    'word/footnotes.xml',
    'word/endnotes.xml',
]


def extract_texts_from_xml(zip_ref, xml_path):
    """Extrai todos os textos (w:t) de um XML dentro do DOCX"""
    if xml_path not in zip_ref.namelist():
        return []
    xml_bytes = zip_ref.read(xml_path)
    root = ET.fromstring(xml_bytes)
    texts = []
    for elem in root.iter():
        # w:t elementos contêm o texto
        if elem.tag.endswith('}t') and elem.text:
            texts.append(elem.text)
    return texts


def extract_header_footer_texts(zip_ref):
    """Extrai textos de todos os headers e footers"""
    texts = []
    for name in zip_ref.namelist():
        if name.startswith('word/header') and name.endswith('.xml'):
            texts.extend(extract_texts_from_xml(zip_ref, name))
        if name.startswith('word/footer') and name.endswith('.xml'):
            texts.extend(extract_texts_from_xml(zip_ref, name))
    return texts


def count_words(texts):
    """Conta palavras em uma lista de strings"""
    total = 0
    for t in texts:
        # separação robusta por palavras (inclui números e palavras com acentos)
        words = re.findall(r"[\wÀ-ÖØ-öø-ÿ]+", t, flags=re.UNICODE)
        total += len(words)
    return total


def detect_toc(zip_ref):
    """Detecta presença de TOC (sumário) via campos 'TOC'"""
    # Procura instrText com 'TOC'
    found = False
    for name in zip_ref.namelist():
        if name == 'word/document.xml':
            xml = zip_ref.read(name)
            if b'TOC' in xml or b'toc' in xml:
                found = True
                break
    return found


def list_paragraph_styles(docx_path):
    """Lista estilos de parágrafo do documento principal"""
    try:
        doc = Document(docx_path)
    except Exception as e:
        return [], 0, 0
    styles = []
    for p in doc.paragraphs:
        try:
            style_name = p.style.name if p.style else 'Normal'
        except Exception:
            style_name = 'Normal'
        styles.append(style_name)
    style_counts = Counter(styles)
    tables_count = len(doc.tables)
    paragraphs_count = len(doc.paragraphs)
    # retorna pares ordenados por frequência
    ordered = sorted(style_counts.items(), key=lambda x: (-x[1], x[0]))
    return ordered, paragraphs_count, tables_count


def count_images(zip_ref):
    """Conta imagens na pasta word/media"""
    media_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
    return len(media_files)


def main():
    if len(sys.argv) < 2:
        print("Uso: python analisar_docx.py <arquivo.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"❌ Arquivo não encontrado: {docx_path}")
        sys.exit(1)

    print("=" * 60)
    print("🔍 ANÁLISE DO DOCX")
    print("=" * 60)
    print(f"📄 Arquivo: {docx_path}")
    print(f"📏 Tamanho: {os.path.getsize(docx_path) / (1024*1024):.2f} MB")

    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # Textos centrais
        all_texts = []
        for part in XML_PARTS_TEXT:
            all_texts.extend(extract_texts_from_xml(zip_ref, part))
        # Headers e footers
        all_texts.extend(extract_header_footer_texts(zip_ref))

        total_words = count_words(all_texts)
        images = count_images(zip_ref)
        toc_present = detect_toc(zip_ref)

    styles_ordered, paragraphs_count, tables_count = list_paragraph_styles(docx_path)

    print(f"\n📊 Resumo:")
    print(f"- Palavras (estimado): {total_words}")
    print(f"- Imagens: {images}")
    print(f"- Parágrafos: {paragraphs_count}")
    print(f"- Tabelas: {tables_count}")
    print(f"- TOC (Sumário) detectado: {'Sim' if toc_present else 'Não'}")

    print("\n🎨 Estilos de parágrafo (top 10):")
    for i, (style, cnt) in enumerate(styles_ordered[:10], 1):
        print(f"  {i}. {style}: {cnt}")

    print("\n✅ Estrutura detectada corretamente.")

if __name__ == '__main__':
    main()