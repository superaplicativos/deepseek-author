#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ANALISADOR DE DOCUMENTO DOCX - LIVRO 3
=====================================
Script para analisar estrutura, conteúdo e elementos visuais do documento
antes da tradução.
"""

import os
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
import re
from collections import Counter

def instalar_dependencias():
    """Instala dependências necessárias"""
    try:
        from docx import Document
    except ImportError:
        print("📦 Instalando python-docx...")
        os.system("pip install python-docx")

def analisar_docx(arquivo_docx):
    """Analisa documento DOCX completo"""
    print(f"🔍 ANALISANDO DOCUMENTO: {arquivo_docx}")
    print("=" * 60)
    
    if not os.path.exists(arquivo_docx):
        print(f"❌ Arquivo não encontrado: {arquivo_docx}")
        return
    
    # Instalar dependências
    instalar_dependencias()
    from docx import Document
    
    # Informações básicas do arquivo
    tamanho_mb = os.path.getsize(arquivo_docx) / (1024 * 1024)
    print(f"📁 Tamanho do arquivo: {tamanho_mb:.2f} MB")
    
    try:
        # Carregar documento
        doc = Document(arquivo_docx)
        
        # Análise de conteúdo textual
        print("\n📝 ANÁLISE DE CONTEÚDO TEXTUAL")
        print("-" * 40)
        
        total_paragrafos = len(doc.paragraphs)
        total_palavras = 0
        total_caracteres = 0
        estilos_encontrados = Counter()
        
        texto_completo = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                palavras = len(para.text.split())
                total_palavras += palavras
                total_caracteres += len(para.text)
                texto_completo += para.text + "\n"
                
                # Contar estilos
                if para.style:
                    estilos_encontrados[para.style.name] += 1
        
        print(f"📊 Total de parágrafos: {total_paragrafos}")
        print(f"📊 Total de palavras: {total_palavras:,}")
        print(f"📊 Total de caracteres: {total_caracteres:,}")
        
        # Análise de estilos
        print(f"\n🎨 ESTILOS DETECTADOS ({len(estilos_encontrados)} tipos)")
        print("-" * 40)
        for estilo, count in estilos_encontrados.most_common():
            print(f"   • {estilo}: {count} parágrafos")
        
        # Análise de tabelas
        print(f"\n📋 TABELAS")
        print("-" * 40)
        total_tabelas = len(doc.tables)
        print(f"📊 Total de tabelas: {total_tabelas}")

        if total_tabelas > 0:
            for i, tabela in enumerate(doc.tables):
                linhas = len(tabela.rows)
                colunas = len(tabela.columns) if tabela.rows else 0
                print(f"   • Tabela {i+1}: {linhas} linhas × {colunas} colunas")

        # Análise de imagens (via ZIP)
        print(f"\n🖼️ ANÁLISE DE IMAGENS")
        print("-" * 40)

        imagens_info = analisar_imagens_zip(arquivo_docx)
        print(f"📊 Total de imagens: {len(imagens_info)}")
        
        if imagens_info:
            tamanho_total_imagens = sum(img['tamanho'] for img in imagens_info)
            print(f"📊 Tamanho total das imagens: {tamanho_total_imagens / (1024*1024):.2f} MB")
            
            # Tipos de imagem
            tipos_imagem = Counter(img['tipo'] for img in imagens_info)
            print(f"📊 Tipos de imagem:")
            for tipo, count in tipos_imagem.items():
                print(f"   • {tipo}: {count} arquivos")
        
        # Análise de hyperlinks e elementos especiais
        print(f"\n🔗 ELEMENTOS ESPECIAIS")
        print("-" * 40)
        
        # Contar hyperlinks no texto
        urls = re.findall(r'https?://\S+|www\.\S+', texto_completo)
        emails = re.findall(r'[\w\.-]+@[\w\.-]+\.[A-Za-z]{2,}', texto_completo)
        
        print(f"📊 URLs encontradas: {len(urls)}")
        print(f"📊 E-mails encontrados: {len(emails)}")
        
        # Análise de formatação especial
        texto_em_caps = re.findall(r'\b[A-Z][A-Z0-9_]{2,}\b', texto_completo)
        texto_entre_crases = re.findall(r'`[^`]+`', texto_completo)
        variaveis_sublinhas = re.findall(r'\b[a-zA-Z_][a-zA-Z0-9_]*_[a-zA-Z0-9_]*\b', texto_completo)

        print(f"📊 Palavras em CAPS (possíveis marcas/nomes): {len(set(texto_em_caps))}")
        print(f"📊 Texto entre crases (código): {len(texto_entre_crases)}")
        print(f"📊 Possíveis variáveis com sublinhado: {len(set(variaveis_sublinhas))}")

        # Elementos detalhados via XML (campos, bookmarks, equações, caixas, hyperlinks, OLE, headers/footers, footnotes/endnotes)
        elementos = coletar_elementos_xml(arquivo_docx)

        print(f"\n📑 LISTA DETALHADA DE ELEMENTOS")
        print("-" * 40)
        print(f"📎 Campos (fldSimple): {elementos.get('fldSimple', 0)}")
        if elementos.get('fields_texts'):
            usados = sorted(set([t.strip() for t in elementos['fields_texts'] if t.strip()]))
            print(f"   • InstrText detectados: {len(usados)} tipos")
            for t in usados[:10]:
                print(f"     - {t[:80]}")
        print(f"🔖 Bookmarks: {elementos.get('bookmarkStart', 0)}")
        print(f"🧮 Equações (OMML): {elementos.get('omath', 0)}")
        print(f"🖼️ Shapes/Desenhos (w:drawing): {elementos.get('drawing', 0)}")
        print(f"🔗 Hyperlinks (w:hyperlink): {elementos.get('hyperlink', 0)}")
        print(f"📦 Objetos OLE/Embeddings: {len(elementos.get('ole_objects', []))}")
        print(f"🧾 Headers: {len(elementos.get('headers', []))} | Footers: {len(elementos.get('footers', []))}")
        print(f"🦶 Footnotes (parágrafos): {elementos.get('footnotes_para', 0)} | Endnotes (parágrafos): {elementos.get('endnotes_para', 0)}")

        # Estimativa de tempo de tradução
        print(f"\n⏱️ ESTIMATIVA DE TRADUÇÃO")
        print("-" * 40)
        
        # Baseado em ~500 palavras por minuto do Google Translate
        tempo_estimado_min = total_palavras / 500
        print(f"⏱️ Tempo estimado por idioma: {tempo_estimado_min:.1f} minutos")
        print(f"⏱️ Tempo total (2 idiomas): {tempo_estimado_min * 2:.1f} minutos")
        
        # Salvar amostra do texto para verificação
        amostra_texto = texto_completo[:1000] + "..." if len(texto_completo) > 1000 else texto_completo
        
        print(f"\n📄 AMOSTRA DO CONTEÚDO (primeiros 1000 caracteres)")
        print("-" * 40)
        print(amostra_texto)
        
        # Verificação de estilos principais e não-padrão
        estilos_principais = {"Title", "Heading 1", "Heading 2", "Heading 3", "Normal", "Body Text", "Caption"}
        estilos_nao_padrao = [e for e in estilos_encontrados.keys() if e not in estilos_principais]

        # Geração de relatório e log
        pasta_saida = Path(arquivo_docx).parent
        relatorio_path = pasta_saida / "RELATORIO_PRECHECK_LIVRO3.txt"
        log_path = pasta_saida / "LOG_OPERACOES_LIVRO3.txt"

        try:
            with open(relatorio_path, 'w', encoding='utf-8') as f:
                f.write("Relatório de Pré-Checagem – Livro 3\n")
                f.write("="*60 + "\n\n")
                f.write(f"Arquivo: {Path(arquivo_docx).name}\n")
                f.write(f"Tamanho: {tamanho_mb:.2f} MB\n\n")

                f.write("1.1 Contagem de palavras\n")
                f.write(f"- Total (estimado): {total_palavras}\n\n")

                f.write("1.2 Imagens\n")
                f.write(f"- Total: {len(imagens_info)}\n")
                if imagens_info:
                    f.write("- Lista (nome, tipo, tamanho bytes):\n")
                    for img in imagens_info[:50]:
                        f.write(f"  • {img['nome']} | {img['tipo']} | {img['tamanho']}\n")
                    if len(imagens_info) > 50:
                        f.write(f"  (+ {len(imagens_info)-50} adicionais)\n")
                f.write("\n")

                f.write("1.3 Elementos detectados\n")
                f.write(f"- Tabelas: {total_tabelas}\n")
                f.write(f"- Shapes/Desenhos: {elementos.get('drawing', 0)}\n")
                f.write(f"- Hyperlinks: {elementos.get('hyperlink', 0)} | URLs detectadas (regex): {len(urls)}\n")
                f.write(f"- E-mails detectados (regex): {len(emails)}\n")
                f.write(f"- Bookmarks: {elementos.get('bookmarkStart', 0)}\n")
                f.write(f"- Equações (OMML): {elementos.get('omath', 0)}\n")
                f.write(f"- Campos: fldSimple={elementos.get('fldSimple', 0)} | instrText={len(elementos.get('fields_texts', []))}\n")
                f.write(f"- TOC (Sumário): {'Sim' if detectar_toc_zip(arquivo_docx) else 'Não'}\n")
                f.write(f"- Headers: {len(elementos.get('headers', []))} | Footers: {len(elementos.get('footers', []))}\n")
                f.write(f"- Footnotes (parágrafos): {elementos.get('footnotes_para', 0)} | Endnotes (parágrafos): {elementos.get('endnotes_para', 0)}\n")
                f.write(f"- Objetos OLE/Embeddings: {len(elementos.get('ole_objects', []))}\n\n")

                f.write("1.4 Verificação de estilos\n")
                f.write("- Estilos principais mais frequentes (top 10):\n")
                for estilo, count in estilos_encontrados.most_common(10):
                    f.write(f"  • {estilo}: {count}\n")
                if estilos_nao_padrao:
                    f.write("- Estilos não-padrão detectados: \n")
                    for e in estilos_nao_padrao[:20]:
                        f.write(f"  • {e}\n")
                    if len(estilos_nao_padrao) > 20:
                        f.write(f"  (+ {len(estilos_nao_padrao)-20} adicionais)\n")
                f.write("\n")

                f.write("1.5 Objetos OLE/embebidos\n")
                if elementos.get('ole_objects'):
                    for o in elementos['ole_objects'][:50]:
                        f.write(f"  • {o}\n")
                    if len(elementos['ole_objects']) > 50:
                        f.write(f"  (+ {len(elementos['ole_objects'])-50} adicionais)\n")
                else:
                    f.write("- Nenhum objeto OLE detectado\n")
                f.write("\n")

                f.write("1.6 Texto não traduzível (para marcar)\n")
                f.write(f"- Texto entre crases: {len(texto_entre_crases)}\n")
                f.write(f"- Possíveis variáveis com sublinhado: {len(set(variaveis_sublinhas))}\n")
                f.write(f"- Palavras em CAPS (possíveis nomes/marcas): {len(set(texto_em_caps))}\n")
                f.write("\n")

                f.write("Resumo adicional\n")
                f.write(f"- Parágrafos: {total_paragrafos}\n")
                f.write(f"- Estilos distintos: {len(estilos_encontrados)}\n")
                f.write(f"- Estimativa de tempo (2 idiomas): {tempo_estimado_min * 2:.1f} min\n")

            # Log de operação
            from datetime import datetime
            with open(log_path, 'a', encoding='utf-8') as lg:
                lg.write(f"[{datetime.now().isoformat()}] PRE-CHECK executado para {Path(arquivo_docx).name}\n")
                lg.write(f"    Palavras={total_palavras} Imagens={len(imagens_info)} Tabelas={total_tabelas} TOC={'Sim' if detectar_toc_zip(arquivo_docx) else 'Não'}\n")
        except Exception:
            pass

        # Resumo final
        print(f"\n✅ RESUMO DA ANÁLISE")
        print("=" * 60)
        print(f"📁 Arquivo: {Path(arquivo_docx).name}")
        print(f"📊 Tamanho: {tamanho_mb:.2f} MB")
        print(f"📝 Palavras: {total_palavras:,}")
        print(f"📋 Tabelas: {total_tabelas}")
        print(f"🖼️ Imagens: {len(imagens_info)}")
        print(f"🎨 Estilos: {len(estilos_encontrados)}")
        print(f"⏱️ Tempo estimado: {tempo_estimado_min * 2:.1f} min")
        print(f"🧾 Relatório salvo em: {relatorio_path}")
        print(f"🗂️ Log atualizado em: {log_path}")

        return {
            'palavras': total_palavras,
            'caracteres': total_caracteres,
            'paragrafos': total_paragrafos,
            'tabelas': total_tabelas,
            'imagens': len(imagens_info),
            'estilos': len(estilos_encontrados),
            'tempo_estimado': tempo_estimado_min * 2
        }
        
    except Exception as e:
        print(f"❌ Erro na análise: {e}")
        return None

def analisar_imagens_zip(arquivo_docx):
    """Analisa imagens dentro do arquivo DOCX via ZIP"""
    imagens = []
    
    try:
        with zipfile.ZipFile(arquivo_docx, 'r') as zip_ref:
            # Listar todos os arquivos
            for arquivo in zip_ref.namelist():
                if arquivo.startswith('word/media/'):
                    info = zip_ref.getinfo(arquivo)
                    nome = Path(arquivo).name
                    extensao = Path(arquivo).suffix.lower()
                    
                    imagens.append({
                        'nome': nome,
                        'caminho': arquivo,
                        'tipo': extensao,
                        'tamanho': info.file_size
                    })
    
    except Exception as e:
        print(f"⚠️ Erro ao analisar imagens: {e}")
    
    return imagens

def detectar_toc_zip(arquivo_docx):
    """Detecta presença de TOC via leitura do ZIP"""
    try:
        with zipfile.ZipFile(arquivo_docx, 'r') as zip_ref:
            if 'word/document.xml' in zip_ref.namelist():
                xml = zip_ref.read('word/document.xml')
                return (b'TOC' in xml or b'toc' in xml)
    except Exception:
        pass
    return False

def coletar_elementos_xml(arquivo_docx):
    """Coleta elementos de interesse a partir dos XMLs do DOCX"""
    elementos = {
        'fldSimple': 0,
        'fields_texts': [],
        'bookmarkStart': 0,
        'omath': 0,
        'drawing': 0,
        'hyperlink': 0,
        'headers': [],
        'footers': [],
        'footnotes_para': 0,
        'endnotes_para': 0,
        'ole_objects': []
    }

    try:
        with zipfile.ZipFile(arquivo_docx, 'r') as zip_ref:
            entries = zip_ref.namelist()

            def parse_xml(name):
                try:
                    xml_bytes = zip_ref.read(name)
                except KeyError:
                    return None
                try:
                    return ET.fromstring(xml_bytes)
                except Exception:
                    return None

            doc_root = parse_xml('word/document.xml')

            if doc_root is not None:
                for el in doc_root.iter():
                    tag = el.tag
                    if tag.endswith('}fldSimple'):
                        elementos['fldSimple'] += 1
                    elif tag.endswith('}instrText'):
                        if el.text:
                            elementos['fields_texts'].append(el.text)
                    elif tag.endswith('}bookmarkStart'):
                        elementos['bookmarkStart'] += 1
                    elif tag.endswith('}oMath') or tag.endswith('}oMathPara'):
                        elementos['omath'] += 1
                    elif tag.endswith('}drawing'):
                        elementos['drawing'] += 1
                    elif tag.endswith('}hyperlink'):
                        elementos['hyperlink'] += 1

            # Headers/Footers
            elementos['headers'] = [f for f in entries if f.startswith('word/header') and f.endswith('.xml')]
            elementos['footers'] = [f for f in entries if f.startswith('word/footer') and f.endswith('.xml')]

            # Footnotes/Endnotes
            foot_root = parse_xml('word/footnotes.xml')
            if foot_root is not None:
                elementos['footnotes_para'] = sum(1 for el in foot_root.iter() if el.tag.endswith('}p'))
            end_root = parse_xml('word/endnotes.xml')
            if end_root is not None:
                elementos['endnotes_para'] = sum(1 for el in end_root.iter() if el.tag.endswith('}p'))

            # OLE/Embeddings
            elementos['ole_objects'] = [f for f in entries if f.startswith('word/embeddings/')]

    except Exception:
        pass

    return elementos

def main():
    """Função principal"""
    if len(sys.argv) < 2:
        print("❌ Uso: python analisar_livro3.py <arquivo.docx>")
        return
    
    arquivo_docx = sys.argv[1]
    resultado = analisar_docx(arquivo_docx)
    
    if resultado:
        print(f"\n🎯 DOCUMENTO PRONTO PARA TRADUÇÃO!")
        print("   ✅ Estrutura detectada corretamente")
        print("   ✅ Imagens e formatação identificadas")
        print("   ✅ Pode prosseguir com a tradução")

if __name__ == "__main__":
    main()